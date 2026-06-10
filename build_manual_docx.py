"""Genera manual_vendedor.docx (Word) a partir del manual, en tono natural."""
from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
AZUL = RGBColor(0x1F, 0x4E, 0x79)


def titulo(texto, size=22):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = AZUL
    p.paragraph_format.space_after = Pt(10)


def seccion(texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = AZUL
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def parrafo(texto):
    p = doc.add_paragraph(texto)
    p.paragraph_format.space_after = Pt(8)
    return p


def parrafo_lead(lead, resto):
    """Párrafo con un arranque en negrita (la categoría) y el resto normal."""
    p = doc.add_paragraph()
    p.add_run(lead).bold = True
    p.add_run(" " + resto)
    p.paragraph_format.space_after = Pt(8)
    return p


def faq(preg, resp):
    p = doc.add_paragraph()
    p.add_run(preg).bold = True
    p.add_run("  " + resp)
    p.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
titulo("Cómo usar a Catu")
parrafo(
    "Catu es el asistente que tenemos en WhatsApp para apoyarte en el día a día. La idea es "
    "simple: en vez de entrar a SAP o llamar a la oficina por cada cosa rápida, le escribes a "
    "Catu y te responde al toque —stock, precios de lista, el estado de tus clientes, sus "
    "pedidos, cobranzas, lo que necesites."
)
parrafo("No es una app ni tiene menús. Le hablas como le hablarías a un compañero por chat.")

seccion("Lo primero")
parrafo(
    "Guarda el número de Catu y mándale un mensaje. No tienes que registrarte ni poner códigos: "
    "te reconoce por tu número y solo trabaja con tu cartera, así que todo lo que veas es tuyo y "
    "de nadie más."
)
parrafo(
    'Si en algún momento la conversación se enreda o quieres empezar de cero, escribe "reiniciar".'
)

seccion("Qué le puedes preguntar")
parrafo(
    "La forma más fácil de entenderlo es con ejemplos. Todo esto funciona tal cual, en lenguaje "
    "normal:"
)
parrafo_lead(
    "Tus clientes y tu cartera.",
    '"¿Qué clientes tengo?", "¿cuántos están activos?", "dame el perfil de Repuestos Razo". No '
    "necesitas el RUC: con el nombre basta, Catu lo busca en tu cartera. Si hay dos parecidos, te "
    "pregunta a cuál te refieres.",
)
parrafo_lead(
    "Crédito y cobranzas.",
    '"¿Cuánto crédito le queda a Transportes Andinos?", "¿tiene deuda vencida?", "¿cuánto me debe '
    'y cuándo vence?", "¿qué clientes míos tienen letras por vencer?".',
)
parrafo_lead(
    "Stock y productos.",
    '"¿Cuántas unidades hay del FIL-BOC-0001?", "búscame filtros de aceite para Toyota", "¿hay '
    'algo equivalente a este código?".',
)
parrafo_lead(
    "Precios.",
    '"¿Cuál es el precio de lista del FIL-BOC-0001?". Ojo con esto: Catu maneja el precio de '
    "lista. Los netos, descuentos por volumen o precios especiales no los da —esos los ves con tu "
    "jefe de línea.",
)
parrafo_lead(
    "Pedidos y entregas.",
    '"¿Cuáles son los últimos pedidos de Transportes Andinos?", "¿en qué estado está el pedido '
    'PED-000123 de ese cliente?", "¿cuándo le llega?". Un consejo: cuando preguntes por un pedido, '
    "dile también de qué cliente es. Así lo encuentra de una.",
)
parrafo_lead(
    "Pagos y documentos.",
    '"¿El pedido tal de Transportes Andinos ya está pagado?", "pásame la factura de ese pedido", '
    '"¿y la guía de remisión?".',
)
parrafo_lead(
    "Vehículos.",
    '"¿Qué repuestos sirven para la placa ABC-123?". También puede consultar la placa en SUNARP '
    "—eso se demora entre 20 y 60 segundos y te llega la foto de la tarjeta.",
)

seccion("Lo que Catu no hace (y por qué)")
parrafo(
    "Conviene tenerlo claro para no perder tiempo. Hay cosas que, a propósito, Catu no resuelve y "
    "te manda al área que corresponde."
)
parrafo(
    "Los precios netos, descuentos por volumen o precios especiales los coordinas con tu jefe de "
    "línea. Una excepción de crédito tampoco la aprueba Catu; eso es de créditos. Si preguntas en "
    "qué almacén está un producto o a qué hora sale el reparto, te va a decir que eso lo ve "
    "logística. Y si un producto está agotado, por ahora no te puede dar la fecha de reposición "
    "—eso lo confirmas con tu jefe de línea."
)
parrafo(
    "Tampoco te va a dar información de un cliente que no es tuyo: si preguntas por uno de otra "
    "cartera, simplemente te dice que no es tuyo."
)
parrafo(
    "Y lo más importante: Catu no inventa. Si no tiene un dato, te lo dice de frente en vez de "
    "adivinar. Si alguna vez te responde algo que te suena raro, confírmalo."
)

seccion("Si un cliente reclama")
parrafo(
    "Cuando un cliente te diga que algo llegó mal —producto equivocado, dañado, lo que sea— "
    'cuéntaselo a Catu con el número de pedido y el motivo. Por ejemplo: "el cliente reclama que '
    'recibió un producto equivocado en el pedido PED-000123". Catu toma los datos y pasa el caso a '
    "atención al cliente para que lo gestionen."
)

seccion("Un par de mañas que ayudan")
parrafo(
    'Háblale natural, no necesitas escribir como robot: "¿cuánto me debe Transportes Andinos?" '
    "funciona igual que cualquier comando rebuscado."
)
parrafo(
    "Nombra al cliente en vez de andar buscando el RUC. Para pedidos y facturas, acompáñalo del "
    "cliente. Y pregunta de a una cosa —si le mandas tres temas mezclados en un solo mensaje, se "
    "complica."
)

seccion("Preguntas que siempre salen")
faq("¿Catu ve los clientes de otros asesores?", "No. Solo el tuyo.")
faq("¿Los precios son los finales?", "Es el precio de lista. El neto lo ves con tu jefe de línea.")
faq("¿Puedo hacer un pedido por Catu?",
    "Por ahora no. Catu es para consultar; el pedido lo registras como siempre.")
faq("¿Por qué a veces se demora?",
    "Casi todo es instantáneo. Lo único lento es la consulta de placa en SUNARP, que puede tomar hasta un minuto.")

doc.save("manual_vendedor.docx")
print("OK -> manual_vendedor.docx")
